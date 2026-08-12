from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Esai_POLITRACK_Literasi_UNEJ.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_centered_paragraph(doc, text="", size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_list_item(doc, text, ordered=False):
    p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_heading(doc, text, level=1, center=False):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_cover(doc):
    add_centered_paragraph(doc, "PERAN WEBSITE POLITRACK SEBAGAI MEDIA", 14, True, 0)
    add_centered_paragraph(doc, "LITERASI POLITIK BAGI PEMILIH MUDA", 14, True, 24)
    add_centered_paragraph(doc, "Disusun untuk memenuhi tugas literasi digital dan demokrasi", 12, False, 36)
    add_centered_paragraph(doc, "Disusun oleh:", 12, False, 8)
    add_centered_paragraph(doc, "Nama: ................................................", 12, False, 0)
    add_centered_paragraph(doc, "NIM/Kelas: ............................................", 12, False, 0)
    add_centered_paragraph(doc, "Program Studi: ........................................", 12, False, 72)
    add_centered_paragraph(doc, "UNIVERSITAS JEMBER", 14, True, 0)
    add_centered_paragraph(doc, "2026", 14, True, 0)
    doc.add_page_break()


def add_approval(doc):
    add_heading(doc, "LEMBAR PENGESAHAN", 1, True)
    add_body(
        doc,
        'Esai berjudul "Peran Website POLITRACK sebagai Media Literasi Politik bagi Pemilih Muda" ini disusun sebagai bentuk kajian terhadap website literasi politik yang dikembangkan dalam lingkungan Literasi UNEJ. Esai ini membahas fungsi website sebagai sarana edukasi demokrasi, verifikasi informasi politik, pengenalan kandidat, pemahaman partai politik, serta penguatan sikap kritis pemilih muda.',
    )
    p = doc.add_paragraph("Jember, ......................... 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(7)
    set_cell_text(table.cell(0, 0), "Mengetahui,", False)
    set_cell_text(table.cell(0, 1), "Penyusun,", False)
    set_cell_text(table.cell(1, 0), "Dosen Pembimbing/Guru Pembimbing", False)
    set_cell_text(table.cell(1, 1), "", False)
    set_cell_text(table.cell(2, 0), "\n\n\n", False)
    set_cell_text(table.cell(2, 1), "\n\n\n", False)
    set_cell_text(table.cell(3, 0), "(........................................)", False)
    set_cell_text(table.cell(3, 1), "(........................................)", False)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "nil")
                tc_borders.append(tag)
            tc_pr.append(tc_borders)
    doc.add_page_break()


def add_main_content(doc):
    add_heading(doc, "PENDAHULUAN", 1, True)
    add_heading(doc, "Latar Belakang", 2)
    for text in [
        "Perkembangan teknologi digital telah mengubah cara masyarakat memperoleh informasi politik. Jika sebelumnya informasi politik lebih banyak diperoleh melalui media massa konvensional, saat ini masyarakat, terutama generasi muda, mendapatkan informasi melalui website, media sosial, mesin pencari, dan platform digital lainnya. Kondisi ini membawa peluang sekaligus tantangan. Di satu sisi, akses informasi menjadi lebih cepat dan luas. Namun di sisi lain, masyarakat juga berhadapan dengan hoaks, disinformasi, klaim politik tanpa data, dan opini yang sering disajikan seolah-olah sebagai fakta.",
        "Dalam konteks demokrasi, kemampuan memahami informasi politik menjadi sangat penting. Pemilih tidak cukup hanya mengetahui nama kandidat atau partai politik, tetapi juga perlu memahami rekam jejak, visi-misi, janji politik, isu strategis, serta sumber informasi yang mendukung sebuah klaim. Tanpa literasi politik yang baik, pemilih mudah dipengaruhi oleh propaganda, narasi emosional, atau informasi yang tidak dapat dipertanggungjawabkan.",
        "Website POLITRACK dalam proyek Literasi UNEJ hadir sebagai salah satu media edukasi demokrasi yang berupaya menjawab persoalan tersebut. Website ini menampilkan informasi mengenai kandidat, partai politik, isu strategis nasional, panduan pemilu, glosarium demokrasi, serta asisten AI berbasis sumber. Dengan pendekatan netral dan berbasis data, POLITRACK dapat menjadi sarana pembelajaran politik yang relevan bagi pemilih muda.",
    ]:
        add_body(doc, text)

    add_heading(doc, "Rumusan Masalah", 2)
    for item in [
        "Bagaimana peran website POLITRACK dalam meningkatkan literasi politik masyarakat?",
        "Fitur apa saja yang terdapat dalam website POLITRACK?",
        "Mengapa verifikasi informasi politik penting bagi pemilih muda?",
        "Bagaimana website POLITRACK dapat mendorong partisipasi demokrasi yang lebih bijak?",
    ]:
        add_list_item(doc, item, ordered=True)

    add_heading(doc, "Tujuan", 2)
    for item in [
        "Menjelaskan peran website POLITRACK sebagai media literasi politik.",
        "Menguraikan fitur-fitur utama yang terdapat dalam website.",
        "Menganalisis pentingnya verifikasi informasi dalam proses demokrasi.",
        "Memberikan gambaran mengenai manfaat website bagi pemilih muda.",
    ]:
        add_list_item(doc, item, ordered=True)

    doc.add_page_break()
    add_heading(doc, "PEMBAHASAN", 1, True)
    sections = [
        ("Gambaran Umum Website POLITRACK", [
            "POLITRACK merupakan platform literasi politik yang mengusung gagasan bahwa pemilih sebaiknya membaca fakta sebelum menentukan pilihan. Pada halaman utama website, POLITRACK memperkenalkan dirinya sebagai platform literasi politik yang berfokus pada agregasi, verifikasi, dan penyajian informasi politik secara netral. Pesan utama yang ditampilkan adalah pentingnya rekam jejak yang dapat ditelusuri, bukan sekadar janji politik, klaim tanpa data, atau retorika.",
            "Website ini tidak dirancang untuk mengarahkan pengguna memilih kandidat atau partai tertentu. Sebaliknya, POLITRACK membantu pengguna memahami informasi politik melalui data yang dapat diperiksa. Hal ini terlihat dari prinsip yang ditampilkan dalam website, yaitu netralitas, sumber tertelusur, dan tanpa rekomendasi pilihan politik. Dengan demikian, posisi website adalah sebagai alat bantu pendidikan politik, bukan alat kampanye.",
        ]),
        ("Fitur Kandidat dan Rekam Jejak", [
            "Salah satu fitur penting dalam POLITRACK adalah direktori kandidat. Melalui fitur ini, pengguna dapat mencari nama kandidat, partai, atau daerah pemilihan. Informasi kandidat yang ditampilkan meliputi biografi, pendidikan, pengalaman jabatan, organisasi, visi-misi, program, klaim, janji politik, garis waktu karier, serta metrik kinerja.",
            "Fitur ini penting karena rekam jejak merupakan salah satu dasar rasional dalam menilai calon pemimpin atau wakil rakyat. Pemilih dapat membandingkan apakah seorang kandidat memiliki pengalaman yang sesuai, apakah janji politiknya pernah direalisasikan, serta apakah klaim yang disampaikan didukung oleh sumber yang jelas. Dengan adanya status verifikasi, pengguna tidak hanya membaca profil, tetapi juga diajak memahami hubungan antara klaim dan bukti.",
        ]),
        ("Fitur Partai Politik", [
            "Selain kandidat, POLITRACK juga menyediakan informasi mengenai partai politik. Website menampilkan profil partai, ideologi, visi, misi, program utama, fokus kebijakan, rekam jejak, hasil pemilu, jumlah kursi DPR, serta informasi organisasi partai. Contoh partai yang tersedia dalam data website antara lain PDI-P, Golkar, Gerindra, PKB, NasDem, PKS, PAN, dan Demokrat.",
            "Informasi partai politik penting karena dalam sistem demokrasi Indonesia, partai memiliki peran besar dalam pencalonan, pembentukan kebijakan, dan kerja parlemen. Pemilih yang memahami karakter partai akan lebih mampu menilai arah kebijakan yang mungkin diperjuangkan oleh kandidat. Dengan demikian, literasi politik tidak berhenti pada figur, tetapi juga mencakup pemahaman terhadap kendaraan politik dan rekam jejak kelembagaan.",
        ]),
        ("Pusat Literasi Politik dan Glosarium Demokrasi", [
            'Website POLITRACK memiliki halaman pusat literasi dan edukasi demokrasi. Di dalamnya terdapat artikel edukasi seperti "Mengenal Sistem Pemilu di Indonesia", "Cara Memeriksa Sumber Informasi Politik", "Memahami Fungsi dan Tugas DPRD", "Membedakan Fakta dan Opini dalam Berita Politik", "Hak Pilih dan Kewajiban Warga Negara", serta "Memahami APBD dan Pentingnya Transparansi Anggaran".',
            "Artikel-artikel tersebut membantu pengguna memahami konsep dasar politik dengan bahasa yang sederhana. Misalnya, artikel tentang pemilu menjelaskan jenis pemilu di Indonesia, sistem proporsional terbuka, daerah pemilihan, dan peran KPU. Artikel tentang DPRD menjelaskan tiga fungsi utama DPRD, yaitu fungsi legislasi, fungsi anggaran, dan fungsi pengawasan. Sementara itu, artikel tentang APBD menjelaskan pentingnya transparansi anggaran untuk mencegah korupsi dan meningkatkan akuntabilitas pejabat publik.",
            "Selain artikel, website juga menyediakan glosarium demokrasi dan tata negara. Di dalamnya terdapat istilah seperti APBD, Bawaslu, Dapil, DPRD, DPT, Fraksi, Hak Interpelasi, KPU, Legislasi, Mahkamah Konstitusi, Perda, RAG, dan Threshold. Glosarium ini bermanfaat bagi pemilih muda yang sering menemui istilah politik, tetapi belum memahami maknanya secara jelas.",
        ]),
        ("Edukasi Anti-Hoaks dan Verifikasi Informasi Politik", [
            "Salah satu kekuatan utama website POLITRACK adalah penekanan pada verifikasi informasi. Dalam era digital, informasi politik dapat menyebar dengan sangat cepat, terutama melalui media sosial dan grup percakapan. Banyak informasi dibuat dengan bahasa emosional, potongan gambar, atau judul provokatif yang belum tentu benar.",
            "Website ini mengajarkan langkah-langkah memeriksa informasi politik, yaitu mengidentifikasi sumber, memeriksa sumber primer, membandingkan dengan sumber lain, memperhatikan konteks, dan mewaspadai tanda bahaya seperti klaim tanpa data atau tangkapan layar tanpa tautan asli. Pendekatan ini penting agar pemilih tidak hanya menjadi penerima informasi, tetapi juga menjadi pemeriksa informasi.",
            'POLITRACK juga memuat prinsip "Pause Dulu" dalam panduan pemilu dan anti-hoaks. Prinsip ini mengajak pengguna berhenti sejenak sebelum membagikan informasi, memeriksa sumber utama, membandingkan beberapa sumber, bertanya kepada pihak yang lebih memahami, dan melaporkan informasi palsu. Sikap seperti ini sangat relevan bagi pemilih muda yang aktif menggunakan media sosial.',
        ]),
        ("Isu Strategis Nasional", [
            "POLITRACK tidak hanya membahas aktor politik, tetapi juga isu-isu strategis nasional. Beberapa isu yang ditampilkan antara lain transisi energi, perlindungan pekerja informal, kebebasan berpendapat, pemberantasan korupsi, pendidikan berkualitas, serta kesehatan dan stunting. Setiap isu dilengkapi ringkasan, deskripsi, konteks, regulasi terkait, dan subisu.",
            "Bagian ini penting karena pemilu seharusnya tidak hanya dipahami sebagai persaingan tokoh, tetapi juga sebagai adu gagasan dan kebijakan. Misalnya, isu transisi energi berkaitan dengan masa depan lingkungan dan pekerja sektor energi. Isu pekerja informal berkaitan dengan perlindungan sosial bagi petani, pedagang kecil, pekerja rumah tangga, dan pekerja gig economy. Isu pemberantasan korupsi berkaitan langsung dengan transparansi anggaran dan kepercayaan publik terhadap pemerintah.",
            "Dengan memahami isu-isu tersebut, pemilih dapat menilai apakah program kandidat dan partai sesuai dengan kebutuhan masyarakat. Pemilih juga dapat melihat bahwa pilihan politik memiliki dampak nyata terhadap pendidikan, kesehatan, lingkungan, ekonomi, dan kebebasan sipil.",
        ]),
        ("Asisten AI Berbasis Sumber", [
            "Website POLITRACK juga menghadirkan fitur asisten AI berbasis RAG atau Retrieval-Augmented Generation. Fitur ini dirancang untuk menjawab pertanyaan pengguna mengenai kandidat dengan memanfaatkan data terverifikasi dan sitasi sumber. Tujuannya adalah mengurangi risiko jawaban yang bersifat spekulatif atau tidak berdasar.",
            "Penggunaan AI dalam literasi politik perlu dilakukan dengan hati-hati. AI tidak boleh menggantikan penilaian kritis manusia, tetapi dapat membantu pengguna menemukan informasi lebih cepat. Jika digunakan dengan sumber yang jelas, fitur ini dapat memperkuat akses informasi. Namun, pengguna tetap perlu membaca sumber, membandingkan informasi, dan tidak menerima jawaban AI secara mentah-mentah.",
        ]),
        ("Manfaat POLITRACK bagi Pemilih Muda", [
            "Pemilih muda merupakan kelompok penting dalam demokrasi karena jumlahnya besar dan memiliki kedekatan dengan teknologi digital. Namun, kedekatan dengan internet tidak otomatis berarti memiliki literasi politik yang baik. Banyak pemilih muda masih rentan terhadap informasi singkat, konten viral, atau narasi yang lebih mengutamakan emosi daripada data.",
            "POLITRACK dapat membantu pemilih muda dengan menyediakan ruang belajar politik yang ringkas, interaktif, dan berbasis sumber. Pengguna dapat mempelajari dasar-dasar pemilu, memahami fungsi lembaga negara, mengenali istilah politik, memeriksa rekam jejak kandidat, membandingkan partai, serta memahami isu kebijakan. Dengan fitur-fitur tersebut, pemilih muda dapat membangun kebiasaan memilih berdasarkan informasi, bukan sekadar popularitas.",
        ]),
        ("Rekomendasi Gambar Pendukung", []),
    ]
    for title, paragraphs in sections:
        add_heading(doc, title, 2)
        for text in paragraphs:
            add_body(doc, text)
    for item in [
        'Screenshot halaman utama POLITRACK. Caption: "Tampilan halaman utama website POLITRACK sebagai platform literasi politik."',
        'Screenshot fitur pencarian kandidat. Caption: "Fitur pencarian kandidat membantu pengguna menelusuri profil dan rekam jejak politik."',
        'Screenshot halaman pusat literasi politik. Caption: "Pusat literasi politik menyediakan artikel edukasi dan glosarium demokrasi."',
        'Screenshot halaman isu strategis nasional. Caption: "Isu strategis membantu pemilih memahami persoalan publik yang perlu dipertimbangkan dalam pemilu."',
        'Screenshot halaman partai politik. Caption: "Informasi partai politik membantu pengguna membandingkan visi, program, dan rekam jejak partai."',
        'Gambar ilustrasi anti-hoaks atau verifikasi informasi. Caption: "Verifikasi informasi penting untuk mencegah penyebaran hoaks politik."',
        'Gambar alur Source-to-Decision Pipeline. Caption: "Alur kerja POLITRACK: dari sumber publik, ekstraksi klaim, verifikasi, pelacakan janji, hingga keputusan pemilih."',
    ]:
        add_list_item(doc, item, ordered=True)

    doc.add_page_break()
    add_heading(doc, "PENUTUP/KESIMPULAN", 1, True)
    for text in [
        "Website POLITRACK dalam proyek Literasi UNEJ merupakan media literasi politik yang relevan bagi masyarakat, terutama pemilih muda. Website ini membantu pengguna memahami politik melalui informasi yang lebih terstruktur, mulai dari profil kandidat, rekam jejak, partai politik, isu strategis, artikel edukasi pemilu, glosarium demokrasi, hingga asisten AI berbasis sumber.",
        "Keunggulan utama POLITRACK terletak pada pendekatannya yang netral, berbasis data, dan menekankan pentingnya verifikasi. Dalam situasi ketika informasi politik sering bercampur dengan opini, propaganda, dan hoaks, platform seperti POLITRACK dapat mendorong pemilih untuk lebih kritis. Pemilih tidak hanya diajak mengetahui siapa yang maju dalam pemilu, tetapi juga memahami apa yang diperjuangkan, bagaimana rekam jejaknya, dan apakah klaimnya dapat dibuktikan.",
        "Dengan demikian, POLITRACK berperan sebagai sarana pendidikan demokrasi digital. Platform ini dapat memperkuat partisipasi politik yang lebih sadar, rasional, dan bertanggung jawab. Demokrasi yang sehat membutuhkan pemilih yang tidak mudah dipengaruhi oleh informasi palsu, tetapi mampu menilai pilihan politik berdasarkan fakta, sumber, dan pertimbangan kepentingan publik.",
    ]:
        add_body(doc, text)

    doc.add_page_break()
    add_heading(doc, "DAFTAR PUSTAKA", 1, True)
    sources = [
        "Badan Pusat Statistik. 2025. Statistik Kesejahteraan Rakyat 2025. Diakses melalui https://www.bps.go.id.",
        "DPR RI. 2024. Profil Anggota DPR RI. Diakses melalui https://www.dpr.go.id/anggota.",
        "Komisi Pemilihan Umum. 2023. Daftar Calon Tetap Pileg 2024. Diakses melalui https://infopemilu.kpu.go.id.",
        "Pemerintah Republik Indonesia. 2017. Undang-Undang Nomor 7 Tahun 2017 tentang Pemilihan Umum.",
        "POLITRACK/Literasi UNEJ. 2026. Website Literasi Politik POLITRACK: Pusat Literasi, Kandidat, Partai, Isu Strategis, dan Asisten AI Berbasis Sumber.",
        'Tim POLITRACK. 2026. "Cara Memeriksa Sumber Informasi Politik." Pusat Literasi Politik.',
        'Tim POLITRACK. 2026. "Hak Pilih dan Kewajiban Warga Negara." Pusat Literasi Politik.',
        'Tim POLITRACK. 2026. "Memahami APBD dan Pentingnya Transparansi Anggaran." Pusat Literasi Politik.',
        'Tim POLITRACK. 2026. "Memahami Fungsi dan Tugas DPRD." Pusat Literasi Politik.',
        'Tim POLITRACK. 2026. "Mengenal Sistem Pemilu di Indonesia." Pusat Literasi Politik.',
    ]
    for source in sources:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(source)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()
    add_heading(doc, "LAMPIRAN", 1, True)
    add_heading(doc, "Lampiran 1. Daftar Fitur Website POLITRACK", 2)
    for item in [
        "Halaman utama platform literasi politik.",
        "Direktori kandidat dan rekam jejak.",
        "Halaman perbandingan kandidat.",
        "Direktori partai politik.",
        "Halaman isu strategis nasional.",
        "Panduan pemilu dan simulasi TPS.",
        "Pusat artikel literasi politik.",
        "Glosarium demokrasi dan tata negara.",
        "Asisten AI berbasis sumber.",
    ]:
        add_list_item(doc, item, ordered=True)
    add_heading(doc, "Lampiran 2. Contoh Istilah dalam Glosarium", 2)
    for item in [
        "APBD: rencana keuangan tahunan pemerintah daerah.",
        "Bawaslu: lembaga pengawas penyelenggaraan pemilu.",
        "Dapil: daerah pemilihan sebagai dasar alokasi kursi wakil rakyat.",
        "DPRD: lembaga perwakilan rakyat di tingkat daerah.",
        "DPT: daftar pemilih tetap yang digunakan dalam pemungutan suara.",
        "KPU: lembaga penyelenggara pemilu.",
        "Perda: peraturan daerah yang dibentuk DPRD bersama kepala daerah.",
        "Threshold: ambang batas suara atau kursi dalam sistem pemilu.",
    ]:
        add_list_item(doc, item, ordered=True)
    add_heading(doc, "Lampiran 3. Catatan Penempatan Gambar", 2)
    add_body(
        doc,
        'Gambar pendukung sebaiknya dimasukkan pada bagian pembahasan setelah subbab yang sesuai. Screenshot halaman utama dapat diletakkan setelah subbab "Gambaran Umum Website POLITRACK". Screenshot fitur kandidat dapat diletakkan setelah subbab "Fitur Kandidat dan Rekam Jejak". Screenshot pusat literasi dapat diletakkan setelah subbab "Pusat Literasi Politik dan Glosarium Demokrasi". Screenshot isu strategis dapat diletakkan setelah subbab "Isu Strategis Nasional".',
    )


def main():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc)
    add_approval(doc)
    add_main_content(doc)
    footer = doc.sections[0].footer
    add_page_number(footer.paragraphs[0])
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
